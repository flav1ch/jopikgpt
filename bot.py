import base64
import json
import logging
import mimetypes
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from dotenv import load_dotenv
from openai import OpenAI
from telegram import Update
from telegram.constants import ChatAction
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    filters,
)

from docx import Document
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader

load_dotenv()

TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
MAX_HISTORY_MESSAGES = int(os.getenv("MAX_HISTORY_MESSAGES", "20"))
MAX_FILE_CHARS = int(os.getenv("MAX_FILE_CHARS", "60000"))
DOWNLOADS_DIR = Path(os.getenv("DOWNLOADS_DIR", "downloads"))
GENERATED_DIR = Path(os.getenv("GENERATED_DIR", "generated"))

SYSTEM_PROMPT = os.getenv(
    "SYSTEM_PROMPT",
    "Ты полезный Telegram-ассистент. Отвечай на русском языке, понятно и по делу.",
)

if not TELEGRAM_TOKEN:
    raise RuntimeError("TELEGRAM_TOKEN не найден. Добавь переменную окружения.")
if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY не найден. Добавь переменную окружения.")

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

client = OpenAI(api_key=OPENAI_API_KEY, base_url=OPENAI_BASE_URL)
DOWNLOADS_DIR.mkdir(exist_ok=True)
GENERATED_DIR.mkdir(exist_ok=True)

TEXT_EXTENSIONS = {
    ".txt", ".md", ".csv", ".json", ".xml", ".html", ".css", ".js",
    ".ts", ".py", ".java", ".php", ".rb", ".go", ".rs", ".c",
    ".cpp", ".h", ".sql", ".log", ".yaml", ".yml",
}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
SUPPORTED_OUTPUT_FORMATS = {"txt", "md", "docx", "xlsx", "csv", "json"}


def get_history(context: ContextTypes.DEFAULT_TYPE) -> List[Dict[str, str]]:
    return context.user_data.get("history", [])


def save_history(context: ContextTypes.DEFAULT_TYPE, history: List[Dict[str, str]]) -> None:
    context.user_data["history"] = history[-MAX_HISTORY_MESSAGES:]


def split_text(text: str, limit: int = 4000) -> List[str]:
    return [text[i : i + limit] for i in range(0, len(text), limit)] or [""]


async def reply_long(update: Update, text: str) -> None:
    for part in split_text(text):
        await update.message.reply_text(part)


def truncate_text(text: str, max_chars: int = MAX_FILE_CHARS) -> str:
    if len(text) <= max_chars:
        return text
    return (
        text[:max_chars]
        + f"\n\n[Файл был обрезан: показаны первые {max_chars} символов из {len(text)}. "
        + "Для больших документов лучше подключить RAG/векторную базу.]"
    )


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        pages.append(f"\n--- Страница {i} ---\n{page_text}")
    return "\n".join(pages).strip()


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_text = []
    for table_index, table in enumerate(doc.tables, start=1):
        table_text.append(f"\n--- Таблица {table_index} ---")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            table_text.append(" | ".join(cells))
    return "\n".join(paragraphs + table_text).strip()


def extract_xlsx(path: Path) -> str:
    workbook = load_workbook(str(path), data_only=True, read_only=True)
    output = []
    for sheet in workbook.worksheets:
        output.append(f"\n--- Лист: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if cell is None else str(cell) for cell in row]
            if any(value.strip() for value in values):
                output.append(" | ".join(values))
    return "\n".join(output).strip()


def extract_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="cp1251", errors="ignore")


def extract_file_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext in {".xlsx", ".xlsm"}:
        return extract_xlsx(path)
    if ext in TEXT_EXTENSIONS:
        return extract_text_file(path)
    raise ValueError(
        "Неподдерживаемый формат. Сейчас поддерживаются: PDF, DOCX, XLSX, TXT, MD, CSV, JSON и обычные текстовые/кодовые файлы."
    )


def image_to_data_url(path: Path) -> str:
    mime_type, _ = mimetypes.guess_type(path.name)
    mime_type = mime_type or "image/jpeg"
    encoded = base64.b64encode(path.read_bytes()).decode("utf-8")
    return f"data:{mime_type};base64,{encoded}"


def ask_model_with_text(user_prompt: str, history: List[Dict[str, str]]) -> str:
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        *history,
        {"role": "user", "content": user_prompt},
    ]
    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=messages,
        temperature=0.7,
    )
    return response.choices[0].message.content or "Не удалось получить ответ."


def ask_model_with_image(prompt: str, image_path: Path) -> str:
    data_url = image_to_data_url(image_path)
    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            },
        ],
        temperature=0.3,
    )
    return response.choices[0].message.content or "Не удалось получить ответ по изображению."


def sanitize_filename(name: str, default: str = "ai_file") -> str:
    name = re.sub(r"[^a-zA-Zа-яА-Я0-9_. -]", "", name).strip()
    name = re.sub(r"\s+", "_", name)
    return name[:80] or default


def strip_code_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        return "\n".join(lines).strip()
    return text


def parse_file_request(text: str) -> tuple[str, str, str]:
    raw = text.strip()
    if raw.startswith("/file"):
        raw = raw.removeprefix("/file").strip()

    output_format = "docx"

    ext_match = re.search(r"\.(txt|md|docx|xlsx|csv|json)\b", raw, re.IGNORECASE)
    if ext_match:
        output_format = ext_match.group(1).lower()
        raw = raw.replace(ext_match.group(0), " ").strip()

    fmt_match = re.search(r"\b(txt|md|docx|xlsx|csv|json)\b", raw, re.IGNORECASE)
    if fmt_match:
        output_format = fmt_match.group(1).lower()
        raw = (raw[: fmt_match.start()] + raw[fmt_match.end() :]).strip()

    if "|" in raw:
        filename_part, prompt = raw.split("|", 1)
        filename = sanitize_filename(filename_part, default=f"ai_file_{datetime.now():%Y%m%d_%H%M%S}")
        prompt = prompt.strip()
    else:
        filename = f"ai_file_{datetime.now():%Y%m%d_%H%M%S}"
        prompt = raw.strip()

    if not prompt:
        prompt = "Создай полезный документ по запросу пользователя."

    return output_format, filename, prompt


def build_generation_prompt(output_format: str, user_prompt: str) -> str:
    if output_format in {"txt", "md"}:
        return (
            "Создай содержимое файла по запросу пользователя. Верни только готовый текст "
            "без пояснений и без markdown-блока.\n\n"
            f"Запрос пользователя:\n{user_prompt}"
        )
    if output_format == "docx":
        return (
            "Создай содержимое DOCX-документа по запросу пользователя. Используй заголовки, "
            "списки и аккуратную структуру. Верни только текст документа без пояснений и без markdown-блока.\n\n"
            f"Запрос пользователя:\n{user_prompt}"
        )
    if output_format == "csv":
        return (
            "Создай CSV-файл по запросу пользователя. Верни только валидный CSV с заголовками, "
            "без пояснений и без markdown-блока.\n\n"
            f"Запрос пользователя:\n{user_prompt}"
        )
    if output_format == "json":
        return (
            "Создай JSON по запросу пользователя. Верни только валидный JSON без пояснений и без markdown-блока.\n\n"
            f"Запрос пользователя:\n{user_prompt}"
        )
    if output_format == "xlsx":
        return (
            "Создай данные для XLSX-таблицы по запросу пользователя. Верни только валидный JSON-массив объектов. "
            "Каждый объект — строка таблицы, ключи — названия колонок. Не добавляй пояснения и markdown-блок.\n\n"
            f"Запрос пользователя:\n{user_prompt}"
        )
    return user_prompt


def save_docx(path: Path, content: str) -> None:
    doc = Document()
    for block in content.split("\n"):
        line = block.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith(("- ", "• ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            doc.add_paragraph(re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


def save_json_file(path: Path, content: str) -> None:
    data = json.loads(strip_code_fences(content))
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def save_xlsx(path: Path, content: str) -> None:
    data = json.loads(strip_code_fences(content))
    if isinstance(data, dict):
        data = data.get("rows") or data.get("data") or [data]
    if not isinstance(data, list):
        raise ValueError("Для XLSX модель должна вернуть JSON-массив объектов.")

    wb = Workbook()
    ws = wb.active
    ws.title = "AI generated"

    if not data:
        ws.append(["Нет данных"])
    elif all(isinstance(row, dict) for row in data):
        headers = list(dict.fromkeys(key for row in data for key in row.keys()))
        ws.append(headers)
        for row in data:
            ws.append([row.get(header, "") for header in headers])
    elif all(isinstance(row, list) for row in data):
        for row in data:
            ws.append(row)
    else:
        ws.append(["Значение"])
        for row in data:
            ws.append([str(row)])

    for column_cells in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = min(max_length + 2, 60)

    wb.save(str(path))


def generate_file_from_prompt(output_format: str, filename: str, user_prompt: str) -> Path:
    if output_format not in SUPPORTED_OUTPUT_FORMATS:
        raise ValueError(f"Формат {output_format} не поддерживается.")

    prompt = build_generation_prompt(output_format, user_prompt)
    content = strip_code_fences(ask_model_with_text(prompt, []))
    filename = sanitize_filename(filename)
    path = GENERATED_DIR / f"{filename}.{output_format}"

    if output_format in {"txt", "md", "csv"}:
        path.write_text(content, encoding="utf-8")
    elif output_format == "docx":
        save_docx(path, content)
    elif output_format == "json":
        save_json_file(path, content)
    elif output_format == "xlsx":
        save_xlsx(path, content)

    return path


def looks_like_file_generation_request(text: str) -> bool:
    lowered = text.lower()
    triggers = ("создай файл", "сгенерируй файл", "сделай файл", "создай документ", "сделай документ")
    formats = (".txt", ".md", ".docx", ".xlsx", ".csv", ".json", " txt", " md", " docx", " xlsx", " csv", " json")
    return any(trigger in lowered for trigger in triggers) and any(fmt in lowered for fmt in formats)


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Привет! *пукнул* Напиши сообщение, загрузи файл или попроси создать файл.\n\n"
        "Команды:\n"
        "/start — сказать здарова\n"
        "/reset — очистить историю диалога (чтобы не спалиться)\n"
        "/file docx Название | запрос — создать файл\n\n"
        "Файлы на вход: PDF, DOCX, XLSX, TXT, CSV, JSON, MD, изображения JPG/PNG/WebP.\n"
        "Файлы на выход: TXT, MD, DOCX, XLSX, CSV, JSON."
    )


async def reset(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    context.user_data["history"] = []
    await update.message.reply_text("История диалога очищена.")


async def file_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not update.message or not update.message.text:
        return

    help_text = (
        "Использование:\n"
        "/file docx Название документа | Напиши коммерческое предложение для юридической фирмы\n"
        "/file xlsx Отчет | Сделай таблицу расходов на месяц\n"
        "/file csv Клиенты | Создай таблицу клиентов с колонками Имя, Телефон, Статус\n\n"
        "Форматы: txt, md, docx, xlsx, csv, json."
    )

    if update.message.text.strip() == "/file":
        await update.message.reply_text(help_text)
        return

    await update.message.chat.send_action(action=ChatAction.UPLOAD_DOCUMENT)
    await update.message.reply_text("Генерирую файл...")

    try:
        output_format, filename, prompt = parse_file_request(update.message.text)
        path = generate_file_from_prompt(output_format, filename, prompt)
        with path.open("rb") as file_obj:
            await update.message.reply_document(document=file_obj, filename=path.name, caption=f"Готово: {path.name}")
    except Exception as exc:
        logger.exception("Ошибка генерации файла")
        await update.message.reply_text(
            "Не удалось создать файл. Попробуй так:\n\n"
            f"{help_text}\n\n"
            f"Техническая ошибка: {exc}"
        )


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not update.message or not update.message.text:
        return

    user_text = update.message.text.strip()
    if not user_text:
        return

    if looks_like_file_generation_request(user_text):
        await update.message.chat.send_action(action=ChatAction.UPLOAD_DOCUMENT)
        await update.message.reply_text("Похоже, ты просишь создать файл. Генерирую...")
        try:
            output_format, filename, prompt = parse_file_request(user_text)
            path = generate_file_from_prompt(output_format, filename, prompt)
            with path.open("rb") as file_obj:
                await update.message.reply_document(document=file_obj, filename=path.name, caption=f"Готово: {path.name}")
            return
        except Exception as exc:
            logger.exception("Ошибка автогенерации файла")
            await update.message.reply_text(
                "Не удалось автоматически создать файл. Лучше используй команду:\n"
                "/file docx Название | твой запрос\n\n"
                f"Техническая ошибка: {exc}"
            )
            return

    await update.message.chat.send_action(action=ChatAction.TYPING)
    history = get_history(context)

    try:
        answer = ask_model_with_text(user_text, history)
        history.append({"role": "user", "content": user_text})
        history.append({"role": "assistant", "content": answer})
        save_history(context, history)
        await reply_long(update, answer)
    except Exception as exc:
        logger.exception("Ошибка при обращении к AI endpoint")
        await update.message.reply_text(
            "Произошла ошибка при обращении к AI endpoint. Проверь OPENAI_BASE_URL, OPENAI_API_KEY и OPENAI_MODEL.\n\n"
            f"Техническая ошибка: {exc}"
        )


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not update.message or not update.message.document:
        return

    document = update.message.document
    file_name = document.file_name or f"file_{document.file_unique_id}"
    safe_name = Path(file_name).name
    file_path = DOWNLOADS_DIR / safe_name
    caption = update.message.caption or "Проанализируй этот файл: сделай краткое содержание и выдели важные моменты."

    await update.message.chat.send_action(action=ChatAction.UPLOAD_DOCUMENT)
    await update.message.reply_text("Файл получен. Скачиваю и читаю содержимое...")

    try:
        tg_file = await context.bot.get_file(document.file_id)
        await tg_file.download_to_drive(file_path)

        ext = file_path.suffix.lower()
        if ext in IMAGE_EXTENSIONS:
            await update.message.chat.send_action(action=ChatAction.TYPING)
            answer = ask_model_with_image(caption, file_path)
            await reply_long(update, answer)
            return

        extracted_text = truncate_text(extract_file_text(file_path))
        if not extracted_text.strip():
            await update.message.reply_text(
                "Я скачал файл, но не смог извлечь из него текст. Если это скан/PDF-картинка, нужен OCR или модель с vision."
            )
            return

        prompt = (
            f"Пользователь загрузил файл: {safe_name}\n"
            f"Инструкция пользователя: {caption}\n\n"
            "Содержимое файла:\n```\n"
            f"{extracted_text}\n```"
        )
        await update.message.chat.send_action(action=ChatAction.TYPING)
        answer = ask_model_with_text(prompt, get_history(context))
        context.user_data["last_file"] = {"name": safe_name, "text": extracted_text}
        await reply_long(update, answer)

    except Exception as exc:
        logger.exception("Ошибка обработки файла")
        await update.message.reply_text(f"Не удалось обработать файл.\n\nТехническая ошибка: {exc}")


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if not update.message or not update.message.photo:
        return

    caption = update.message.caption or "Опиши изображение и выдели важные детали."
    photo = update.message.photo[-1]
    file_path = DOWNLOADS_DIR / f"photo_{photo.file_unique_id}.jpg"

    await update.message.reply_text("Фото получено. Анализирую...")

    try:
        tg_file = await context.bot.get_file(photo.file_id)
        await tg_file.download_to_drive(file_path)
        answer = ask_model_with_image(caption, file_path)
        await reply_long(update, answer)
    except Exception as exc:
        logger.exception("Ошибка анализа изображения")
        await update.message.reply_text(
            "Не удалось проанализировать изображение. Возможно, твой endpoint не поддерживает vision.\n\n"
            f"Техническая ошибка: {exc}"
        )


async def handle_non_supported(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Я пока принимаю текст, документы и изображения. Для аудио/видео нужно отдельно подключить распознавание."
    )


def main() -> None:
    app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("reset", reset))
    app.add_handler(CommandHandler("file", file_command))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(~filters.COMMAND, handle_non_supported))

    logger.info("Бот запущен через polling")
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
